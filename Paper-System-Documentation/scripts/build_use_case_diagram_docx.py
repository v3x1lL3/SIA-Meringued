#!/usr/bin/env python3
"""Build 02-Use-Case-Diagram.docx with detailed UML-style figures (SVG → PNG)."""

from __future__ import annotations

import io
import sys
from pathlib import Path

SCRIPTS = Path(__file__).resolve().parent
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

try:
    import cairosvg
except ImportError as exc:  # pragma: no cover
    print("Missing cairosvg. Install with: pip install cairosvg", file=sys.stderr)
    raise SystemExit(1) from exc

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from use_case_svg_render import Dependency, render_figure


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
OUT_DOCX = ROOT / "02-Use-Case-Diagram.docx"
OUT_MD = ROOT / "04-Use-Case-Diagrams-Context.md"

# Parallel to FIGURES order: mapping to `02-Use-Case-List.md` / `01-Events-Table.md` sections.
FIGURE_LIST_REFS: list[str] = [
    "A.1 Access and Session — aligns with events table §A",
    "A.2 Dashboard and Order Visibility — aligns with events table §B",
    "A.3 POS Configuration (pricing & fulfillment slice) — aligns with events table §C",
    "A.3 POS Configuration (design, date, media slice) — aligns with events table §C",
    "A.4 Cart, Checkout, and Proof of Payment — aligns with events table §C",
    "A.5 Logs, Receipts, and Exports — aligns with events table §D",
    "B.1–B.2 Admin Access and Dashboard — aligns with events table §A & §E (admin slice)",
    "B.3 Orders Management — aligns with events table §E",
    "B.4 Inventory Management — aligns with events table §F (inventory slice)",
    "B.5 Records Management — aligns with events table §F (records slice)",
]


def _rng(n: int) -> list[int]:
    return list(range(n))


def svg_to_png_bytes(svg: str, scale: float = 1.85) -> bytes:
    return cairosvg.svg2png(bytestring=svg.encode("utf-8"), scale=scale)


FIGURES: list[dict[str, object]] = [
    {
        "file": "use-case-c01-customer-access-session.svg",
        "heading": "Figure 1 — Customer: Access & Session",
        "title": "Customer — Access & Session",
        "actor": "Customer",
        "primaries": [
            "View Landing & Navigate Modules",
            "Start Customer Login / Signup Flow",
            "Register Customer Account",
            "Authenticate & Open Dashboard",
            "Use Protected Customer Routes",
            "Logout & Terminate Session",
        ],
        "secondaries": [
            "Enforce Role-Based Route Access",
            "Create Customer Session State",
        ],
        "actor_to": _rng(6),
        "deps": [
            Dependency(3, 1, "<<include>>"),
            Dependency(4, 0, "<<include>>"),
        ],
    },
    {
        "file": "use-case-c02-customer-dashboard-orders.svg",
        "heading": "Figure 2 — Customer: Dashboard & Order Visibility",
        "title": "Customer — Dashboard & Orders",
        "actor": "Customer",
        "primaries": [
            "Load Dashboard Summary",
            "View Active Order Listing",
            "Track Pipeline Status Changes",
            "Open Order Details Modal",
            "Apply Customer Order Pagination",
        ],
        "secondaries": [
            "Merge / Refresh Customer Order Data",
            "Poll Status While Page Visible",
        ],
        "actor_to": _rng(5),
        "deps": [
            Dependency(1, 0, "<<include>>"),
            Dependency(2, 1, "<<extend>>"),
        ],
    },
    {
        "file": "use-case-c03-customer-pos-pricing.svg",
        "heading": "Figure 3 — Customer: POS Configuration (Pricing & Fulfillment)",
        "title": "Customer — POS: Pricing & Fulfillment",
        "actor": "Customer",
        "primaries": [
            "Open POS Modal & Initialize Defaults",
            "Select Cake Flavor & Dependent Options",
            "Select Frosting & Small/Med/Large Tier",
            "Update Quantity & Recalculate Totals",
            "Configure Payment Plan (Full or 50%)",
            "Configure Pickup or Delivery Mode",
        ],
        "secondaries": [
            "Rebuild Frosting Options by Flavor",
            "Apply Size Tier to Unit Pricing",
        ],
        "actor_to": _rng(6),
        "deps": [
            Dependency(1, 0, "<<include>>"),
            Dependency(2, 1, "<<include>>"),
        ],
    },
    {
        "file": "use-case-c04-customer-pos-design.svg",
        "heading": "Figure 4 — Customer: POS Design, Date & References",
        "title": "Customer — POS: Design & Media",
        "actor": "Customer",
        "primaries": [
            "Set Required Consumption Date",
            "Capture Dedication & Design Instructions",
            "Upload Multiple Design Reference Images",
            "Remove One Image or Clear All Uploads",
        ],
        "secondaries": [
            "Validate File Types & Cumulative Size",
        ],
        "actor_to": _rng(4),
        "deps": [
            Dependency(2, 0, "<<include>>"),
        ],
    },
    {
        "file": "use-case-c05-customer-cart-checkout.svg",
        "heading": "Figure 5 — Customer: Cart, Checkout & Confirmation",
        "title": "Customer — Cart & Checkout",
        "actor": "Customer",
        "primaries": [
            "Add Configured Order to Cart",
            "Review Cart Lines & Computed Totals",
            "Place Order from Cart or Directly",
            "Upload Payment Receipt File",
            "Confirm Order & Persist Order Record",
        ],
        "secondaries": [
            "Validate Checkout & Attachment Rules",
            "Bind Receipt to Target Order",
        ],
        "actor_to": _rng(5),
        "deps": [
            Dependency(2, 0, "<<include>>"),
            Dependency(3, 1, "<<include>>"),
            Dependency(4, 0, "<<include>>"),
        ],
    },
    {
        "file": "use-case-c06-customer-exports.svg",
        "heading": "Figure 6 — Customer: Receipts & PDF Export",
        "title": "Customer — Receipts & PDFs",
        "actor": "Customer",
        "primaries": [
            "Preview Uploaded Receipt",
            "Download Single-Order Printable PDF",
            "Download All-Orders PDF with Letterhead",
        ],
        "secondaries": [
            "Render Print-Ready PDF Layout",
        ],
        "actor_to": _rng(3),
        "deps": [
            Dependency(1, 0, "<<include>>"),
            Dependency(2, 0, "<<include>>"),
        ],
    },
    {
        "file": "use-case-a01-admin-access-dashboard.svg",
        "heading": "Figure 7 — Admin/Staff: Access, Session & Dashboard",
        "title": "Admin — Access & Dashboard",
        "actor": "Admin/Staff",
        "primaries": [
            "Open Admin Authentication Page",
            "Register Admin Account",
            "Authenticate Admin Session",
            "Access Admin Modules by Role",
            "Logout Admin Session",
            "Load Dashboard Metrics & Alerts",
            "Navigate Orders / Inventory / Records",
        ],
        "secondaries": [
            "Create Admin Session with Role Check",
            "Show Pending Orders & Low-Stock Alerts",
        ],
        "actor_to": _rng(7),
        "deps": [
            Dependency(2, 0, "<<include>>"),
            Dependency(5, 1, "<<include>>"),
        ],
    },
    {
        "file": "use-case-a02-admin-orders.svg",
        "heading": "Figure 8 — Admin/Staff: Orders Management",
        "title": "Admin — Orders Management",
        "actor": "Admin/Staff",
        "primaries": [
            "Load Merged Admin Order List",
            "Filter Orders by Status Tabs",
            "Paginate Admin Orders List",
            "View Order Details & Design Media",
            "Verify Customer Payment Receipt",
            "Advance Order Through Pipeline Stages",
            "Manual Status Update or Cancellation",
            "Execute Baking Misc Stock-Out",
        ],
        "secondaries": [
            "Merge & Deduplicate Cloud / Local Sources",
            "Persist Status & Inventory Side Effects",
        ],
        "actor_to": _rng(8),
        "deps": [
            Dependency(0, 0, "<<include>>"),
            Dependency(5, 1, "<<include>>"),
            Dependency(7, 1, "<<include>>"),
        ],
    },
    {
        "file": "use-case-a03-admin-inventory.svg",
        "heading": "Figure 9 — Admin/Staff: Inventory Management",
        "title": "Admin — Inventory Management",
        "actor": "Admin/Staff",
        "primaries": [
            "Switch Ingredients / Miscellaneous Tab",
            "Add New Inventory Item",
            "Search Inventory by Item Name",
            "Paginate Inventory List",
            "Edit Item Metadata & Image",
            "Save Unit Price for Costing",
            "Record Stock-In Movement",
            "Record Stock-Out Movement",
            "View Low-Stock & Expiry Notices",
            "Delete Inventory Item",
        ],
        "secondaries": [
            "Validate Item Fields on Save",
            "Update Quantity & Movement Log",
        ],
        "actor_to": _rng(10),
        "deps": [
            Dependency(1, 0, "<<include>>"),
            Dependency(4, 0, "<<include>>"),
            Dependency(6, 1, "<<include>>"),
            Dependency(7, 1, "<<include>>"),
        ],
    },
    {
        "file": "use-case-a04-admin-records.svg",
        "heading": "Figure 10 — Admin/Staff: Records Management",
        "title": "Admin — Records Management",
        "actor": "Admin/Staff",
        "primaries": [
            "Load Records Module by Context",
            "Switch Record Type Tabs",
            "Apply Date Preset or Custom Range",
            "Search Title / Reference / Notes",
            "Paginate Records List",
            "Add or Edit Allowed Record Types",
            "View Locked Audit / EOD Entries",
            "View Daily Totals Where Applicable",
            "Export Active Record Type to PDF",
        ],
        "secondaries": [
            "Refetch Dataset for Type & Filters",
            "Apply Letterhead to Print Export",
        ],
        "actor_to": _rng(9),
        "deps": [
            Dependency(1, 0, "<<include>>"),
            Dependency(8, 1, "<<include>>"),
        ],
    },
]


def write_svgs() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    for spec in FIGURES:
        svg = render_figure(
            title=str(spec["title"]),
            actor=str(spec["actor"]),
            primaries=list(spec["primaries"]),  # type: ignore[arg-type]
            secondaries=list(spec["secondaries"]),  # type: ignore[arg-type]
            actor_to_primaries=list(spec["actor_to"]),  # type: ignore[arg-type]
            dependencies=list(spec["deps"]),  # type: ignore[arg-type]
        )
        path = ASSETS / str(spec["file"])
        path.write_text(svg, encoding="utf-8")


def write_context_markdown() -> None:
    """Emit companion markdown describing each UML figure (kept in sync with FIGURES)."""
    if len(FIGURE_LIST_REFS) != len(FIGURES):
        raise RuntimeError("FIGURE_LIST_REFS length must match FIGURES")

    lines: list[str] = [
        "# Use Case Diagrams — Textual Context",
        "",
        "> **Regeneration:** Running `scripts/build_use_case_diagram_docx.py` rewrites this file from the same `FIGURES` specification used for the SVG assets and Word export, so the narrative stays aligned with the diagrams.",
        "",
        "This document supplements the graphical UML-style figures embedded in **`02-Use-Case-Diagram.docx`** (source SVGs under **`assets/`**). For tabular and narrative requirements, cross-check **`01-Events-Table.md`**, **`02-Use-Case-List.md`**, and the detailed flowchart in **`03-Use-Case-Diagram.md`**. For paragraph-style module narratives (actors, behaviors, includes/extends, purpose), see **`05-Use-Case-Descriptions.md`**.",
        "",
        "## How to read connections (all figures)",
        "",
        "| Drawing | Meaning |",
        "|---------|---------|",
        "| **Solid line** from the stick figure to an oval | **Association:** the actor participates in that primary use case. |",
        "| **Dashed line** between ovals, labeled `<<include>>` | **Include:** the base use case (usually left) **always** brings in the supporting behavior (right). |",
        "| **Dashed line** between ovals, labeled `<<extend>>` | **Extend:** the extension use case (right) **may** add behavior under conditions; the diagram still draws a dashed link from the **primary** oval toward the **supporting** oval to match the inventory reference layout. |",
        "",
        "Below, each figure spells out **every** actor→use case link and **every** dashed dependency so you can trace the drawing line-by-line.",
        "",
        "## Figure index",
        "",
        "| # | Word / SVG focus | List / events alignment |",
        "|---|------------------|-------------------------|",
    ]
    for i, spec in enumerate(FIGURES, start=1):
        lines.append(
            f"| {i} | {spec['heading']} | {FIGURE_LIST_REFS[i - 1]} |"
        )
    lines.extend(["", "---", ""])

    for i, spec in enumerate(FIGURES, start=1):
        primaries: list[str] = list(spec["primaries"])  # type: ignore[arg-type]
        secondaries: list[str] = list(spec["secondaries"])  # type: ignore[arg-type]
        deps: list[Dependency] = list(spec["deps"])  # type: ignore[arg-type]
        actor_to: list[int] = list(spec["actor_to"])  # type: ignore[arg-type]
        actor = str(spec["actor"])

        lines.extend(
            [
                f"## {spec['heading']}",
                "",
                f"- **List / events context:** {FIGURE_LIST_REFS[i - 1]}",
                f"- **System boundary label (diagram title bar):** `{spec['title']}`",
                f"- **Actor:** `{actor}`",
                f"- **Asset file:** `assets/{spec['file']}`",
                "",
                "### Connection guide (what is wired to what)",
                "",
                "#### A. Actor → primary use case (**solid lines**)",
                "",
            ]
        )
        if not actor_to:
            lines.append("_No actor→primary links defined for this figure._")
        else:
            for pi in actor_to:
                pname = primaries[pi]
                lines.append(
                    f"- **`{actor}`** —(solid line)→ **`{pname}`**"
                )

        lines.extend(["", "#### B. Primary → supporting (**dashed** `<<include>>` / `<<extend>>`)", ""])
        if not deps:
            lines.append("_No dashed dependencies in this figure._")
        else:
            for dep in deps:
                p_name = primaries[dep.primary_index]
                s_name = secondaries[dep.secondary_index]
                stereotype = dep.stereotype
                lines.append(
                    f"- **`{p_name}`** —(dashed, {stereotype})→ **`{s_name}`**"
                )

        outgoing = {d.primary_index for d in deps}
        incoming_sec = {d.secondary_index for d in deps}
        actor_only = [primaries[j] for j in range(len(primaries)) if j not in outgoing]
        if actor_only:
            lines.extend(
                [
                    "",
                    "#### C. Primary use cases with **only** the actor link (no dashed stereotype in this figure)",
                    "",
                    "These ovals still appear in the diagram and are reached by the stick figure, but they are **not** the source of an `<<include>>` or `<<extend>>` arrow in this particular figure:",
                    "",
                ]
            )
            for name in actor_only:
                lines.append(f"- `{name}`")
        unused_sec = [
            secondaries[j] for j in range(len(secondaries)) if j not in incoming_sec
        ]
        if unused_sec:
            lines.extend(
                [
                    "",
                    "#### Supporting ovals not used as a dependency target (check spec)",
                    "",
                ]
            )
            for name in unused_sec:
                lines.append(f"- `{name}`")

        lines.extend(
            [
                "",
                "### Primary use cases (reference list)",
                "",
            ]
        )
        for j, name in enumerate(primaries, start=1):
            lines.append(f"{j}. {name}")
        lines.extend(["", "### Supporting use cases (reference list)", ""])
        for j, name in enumerate(secondaries, start=1):
            lines.append(f"{j}. {name}")
        lines.append("")

    lines.extend(
        [
            "## Shared / system behavior",
            "",
            "Cross-cutting behaviors from **`02-Use-Case-List.md` § C** (route guards, merge/dedupe, payload limits, pagination state, validation) appear partly as supporting use cases on the diagrams and partly as implementation detail behind the events in **`01-Events-Table.md` § G**.",
            "",
        ]
    )

    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    write_svgs()
    write_context_markdown()

    doc = Document()
    title = doc.add_heading("Use Case Diagram", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "These figures follow the same UML-style illustration pattern as the inventory reference: "
        "actors on the left, a large system boundary oval, teal primary use case ovals, and dashed "
        "<<include>> / <<extend>> dependencies to supporting use cases on the right. "
        "Labels are expanded to mirror the paper’s use case list and events table (group A–G), "
        "split across multiple diagrams so the layout stays readable."
    )
    for run in intro.runs:
        run.font.size = Pt(11)

    for spec in FIGURES:
        doc.add_heading(str(spec["heading"]), level=1)
        svg = render_figure(
            title=str(spec["title"]),
            actor=str(spec["actor"]),
            primaries=list(spec["primaries"]),  # type: ignore[arg-type]
            secondaries=list(spec["secondaries"]),  # type: ignore[arg-type]
            actor_to_primaries=list(spec["actor_to"]),  # type: ignore[arg-type]
            dependencies=list(spec["deps"]),  # type: ignore[arg-type]
        )
        doc.add_picture(io.BytesIO(svg_to_png_bytes(svg)), width=Inches(6.45))

    note = doc.add_paragraph()
    note.add_run("Actor naming (").bold = True
    note.add_run("Admin/Staff")
    note.add_run("): ").bold = True
    note.add_run(
        "The back office is drawn as a single actor labeled Admin/Staff to reflect that operational "
        "employees share one administrative application surface, while the events table still refers "
        "to an Admin user as the authenticated party. If your instructor prefers strict UML, you may "
        "rename the actor to Admin only; the underlying behavior is unchanged."
    )

    doc.save(OUT_DOCX)
    print(f"Wrote {len(FIGURES)} SVGs under {ASSETS}")
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
